#!/usr/bin/env python3
"""
Brochure Link Extractor — French Wedding Style
Downloads files from Google Drive and extracts text content.

Usage:
  python extract_brochure.py <url_type> <url_or_id> <access_token> <working_dir> [<folder_id>]

URL types:
  GDRIVE_FILE   — Single Google Drive file (PDF/DOCX). Pass the file URL.
  GDRIVE_FOLDER — Google Drive folder. Pass folder ID. Lists + downloads all files.
  GDOCS         — Google Doc. Pass document ID. Exports as plain text.

Output:
  On success: prints extracted markdown to stdout
  On failure: prints error line starting with [ERROR]: to stdout
  Progress/debug messages go to stderr

Dependencies: pypdf, python-docx (install via: python -m pip install pypdf python-docx)
"""

import sys
import os
import re
import json
import urllib.request
import urllib.error
import urllib.parse
import tempfile
import io

# Force UTF-8 for stdout/stderr on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')


def log(msg):
    """Print to stderr for progress messages."""
    print(msg, file=sys.stderr)


def extract_file_id_from_url(url):
    """Extract Google Drive file ID from various URL formats."""
    # drive.google.com/file/d/{ID}/...
    m = re.search(r'/file/d/([a-zA-Z0-9_-]+)', url)
    if m:
        return m.group(1)
    # drive.google.com/open?id={ID}
    m = re.search(r'[?&]id=([a-zA-Z0-9_-]+)', url)
    if m:
        return m.group(1)
    return None


def extract_folder_id_from_url(url):
    """Extract Google Drive folder ID from URL."""
    m = re.search(r'/folders/([a-zA-Z0-9_-]+)', url)
    if m:
        return m.group(1)
    return None


def extract_doc_id_from_url(url):
    """Extract Google Docs document ID from URL."""
    m = re.search(r'/document/d/([a-zA-Z0-9_-]+)', url)
    if m:
        return m.group(1)
    return None


def gdrive_download(file_id, access_token, dest_path):
    """Download a file from Google Drive using the API."""
    url = f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {access_token}"
    })
    try:
        with urllib.request.urlopen(req) as resp:
            with open(dest_path, 'wb') as f:
                f.write(resp.read())
        return True
    except urllib.error.HTTPError as e:
        log(f"  Download failed: HTTP {e.code}")
        return False


def gdrive_get_file_meta(file_id, access_token):
    """Get file metadata (name, mimeType) from Google Drive."""
    url = f"https://www.googleapis.com/drive/v3/files/{file_id}?fields=name,mimeType,size"
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {access_token}"
    })
    try:
        with urllib.request.urlopen(req) as resp:
            return json.loads(resp.read())
    except urllib.error.HTTPError:
        return None


def gdrive_list_folder(folder_id, access_token):
    """List files in a Google Drive folder."""
    url = (f"https://www.googleapis.com/drive/v3/files"
           f"?q=%27{folder_id}%27+in+parents+and+trashed%3Dfalse"
           f"&fields=files(id,name,mimeType,size)"
           f"&pageSize=100")
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {access_token}"
    })
    try:
        with urllib.request.urlopen(req) as resp:
            data = json.loads(resp.read())
            return data.get('files', [])
    except urllib.error.HTTPError as e:
        log(f"  Folder listing failed: HTTP {e.code}")
        return None


def gdrive_export_doc(doc_id, access_token):
    """Export a Google Doc as plain text."""
    url = f"https://www.googleapis.com/drive/v3/files/{doc_id}/export?mimeType=text/plain"
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {access_token}"
    })
    try:
        with urllib.request.urlopen(req) as resp:
            return resp.read().decode('utf-8', errors='replace')
    except urllib.error.HTTPError as e:
        log(f"  Doc export failed: HTTP {e.code}")
        return None


def clean_text(text):
    """Clean extracted text — fix encoding artifacts, collapse whitespace."""
    # Replace replacement character with a space
    text = text.replace('\ufffd', ' ')
    # Fix common PDF extraction artifacts: spaced-out letters
    # e.g., "E v e n t  B r o c h u r e" → keep as-is (title styling)
    # Collapse multiple spaces within lines (but not all — some are intentional)
    text = re.sub(r'[ \t]{3,}', '  ', text)
    # Collapse 3+ consecutive blank lines to 2
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    return text.strip()


def extract_pdf_text(filepath):
    """Extract text from a PDF file using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(filepath)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(clean_text(text))
        return '\n\n'.join(pages)
    except Exception as e:
        log(f"  PDF extraction failed: {e}")
        return None


def extract_docx_text(filepath):
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Also extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(' | '.join(cells))
            if rows:
                paragraphs.append('\n'.join(rows))

        return '\n\n'.join(paragraphs)
    except Exception as e:
        log(f"  DOCX extraction failed: {e}")
        return None


def extract_text_from_file(filepath, filename, mime_type=None):
    """Extract text based on file type."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == '.pdf' or (mime_type and 'pdf' in mime_type):
        return extract_pdf_text(filepath)
    elif ext in ('.docx',) or (mime_type and 'wordprocessing' in str(mime_type)):
        return extract_docx_text(filepath)
    elif ext in ('.txt', '.md', '.csv'):
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    elif ext in ('.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.svg'):
        return None  # Images — no OCR available
    else:
        return None  # Unsupported format


def is_google_workspace_type(mime_type):
    """Check if the file is a Google Workspace type (Docs, Sheets, Slides)."""
    workspace_types = [
        'application/vnd.google-apps.document',
        'application/vnd.google-apps.spreadsheet',
        'application/vnd.google-apps.presentation',
    ]
    return mime_type in workspace_types


def process_gdrive_file(url, access_token, working_dir):
    """Process a single Google Drive file URL."""
    file_id = extract_file_id_from_url(url)
    if not file_id:
        return "[ERROR]: Could not extract file ID from URL."

    log(f"  File ID: {file_id}")

    # Get file metadata
    meta = gdrive_get_file_meta(file_id, access_token)
    if not meta:
        return "[ERROR]: Link unreachable."

    filename = meta.get('name', 'unknown')
    mime_type = meta.get('mimeType', '')
    log(f"  File: {filename} ({mime_type})")

    # Handle Google Workspace files via export
    if is_google_workspace_type(mime_type):
        if 'document' in mime_type:
            text = gdrive_export_doc(file_id, access_token)
            if text and text.strip():
                return text.strip()
            return "[ERROR]: Files are images without readable text."
        elif 'spreadsheet' in mime_type:
            # Export as CSV
            export_url = f"https://www.googleapis.com/drive/v3/files/{file_id}/export?mimeType=text/csv"
            req = urllib.request.Request(export_url, headers={
                "Authorization": f"Bearer {access_token}"
            })
            try:
                with urllib.request.urlopen(req) as resp:
                    text = resp.read().decode('utf-8', errors='replace')
                    if text.strip():
                        return text.strip()
            except Exception:
                pass
            return "[ERROR]: Files are images without readable text."
        else:
            return "[ERROR]: Unsupported file types found."

    # Download binary file
    dest = os.path.join(working_dir, f"temp_{file_id}{os.path.splitext(filename)[1]}")
    if not gdrive_download(file_id, access_token, dest):
        return "[ERROR]: Link unreachable."

    try:
        text = extract_text_from_file(dest, filename, mime_type)
        if text and text.strip():
            return text.strip()
        else:
            return "[ERROR]: Files are images without readable text."
    finally:
        if os.path.exists(dest):
            os.remove(dest)


def process_gdrive_folder(folder_id, access_token, working_dir):
    """Process a Google Drive folder — list and extract all files."""
    log(f"  Folder ID: {folder_id}")
    files = gdrive_list_folder(folder_id, access_token)

    if files is None:
        return "[ERROR]: Link unreachable."
    if len(files) == 0:
        return "[ERROR]: Link unreachable."

    log(f"  Found {len(files)} files in folder")

    results = []
    errors = []

    for f in files:
        fid = f['id']
        fname = f['name']
        fmime = f.get('mimeType', '')
        log(f"  Processing: {fname} ({fmime})")

        # Skip folders
        if fmime == 'application/vnd.google-apps.folder':
            log(f"    Skipping subfolder")
            continue

        text = None

        # Handle Google Workspace files
        if is_google_workspace_type(fmime):
            if 'document' in fmime:
                text = gdrive_export_doc(fid, access_token)
            elif 'spreadsheet' in fmime:
                export_url = f"https://www.googleapis.com/drive/v3/files/{fid}/export?mimeType=text/csv"
                req = urllib.request.Request(export_url, headers={
                    "Authorization": f"Bearer {access_token}"
                })
                try:
                    with urllib.request.urlopen(req) as resp:
                        text = resp.read().decode('utf-8', errors='replace')
                except Exception:
                    pass
            # Slides — skip for now
        else:
            # Download and extract
            ext = os.path.splitext(fname)[1].lower()
            if ext in ('.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.svg'):
                log(f"    Image file — no OCR available, skipping")
                errors.append(fname)
                continue
            if ext in ('.zip', '.exe', '.rar', '.7z'):
                log(f"    Unsupported format, skipping")
                errors.append(fname)
                continue

            dest = os.path.join(working_dir, f"temp_{fid}{ext}")
            if gdrive_download(fid, access_token, dest):
                try:
                    text = extract_text_from_file(dest, fname, fmime)
                finally:
                    if os.path.exists(dest):
                        os.remove(dest)

        if text and text.strip():
            results.append((fname, text.strip()))
        else:
            errors.append(fname)

    if not results:
        if errors:
            return "[ERROR]: Files are images without readable text."
        return "[ERROR]: Link unreachable."

    # Consolidate with headers + dividers
    parts = []
    if errors and results:
        parts.append(f"[WARNING]: Some files in folder were unreadable ({', '.join(errors)}).\n")

    for i, (fname, content) in enumerate(results):
        parts.append(f"### Source: {fname}\n\n{content}")
        if i < len(results) - 1:
            parts.append("---")

    return '\n\n'.join(parts)


def process_gdocs(doc_id, access_token, working_dir):
    """Process a Google Doc — export as plain text, or download if it's an uploaded file."""
    log(f"  Doc ID: {doc_id}")

    # First check the actual mimeType
    meta = gdrive_get_file_meta(doc_id, access_token)
    if not meta:
        return "[ERROR]: Link unreachable."

    filename = meta.get('name', 'unknown')
    mime_type = meta.get('mimeType', '')
    log(f"  File: {filename} ({mime_type})")

    # Native Google Doc — use export
    if is_google_workspace_type(mime_type):
        text = gdrive_export_doc(doc_id, access_token)
        if text and text.strip():
            return text.strip()
        return "[ERROR]: Files are images without readable text."

    # Uploaded file (e.g., .docx opened in Google Docs) — download binary
    ext = os.path.splitext(filename)[1].lower()
    dest = os.path.join(working_dir, f"temp_{doc_id}{ext}")
    if not gdrive_download(doc_id, access_token, dest):
        return "[ERROR]: Link unreachable."

    try:
        text = extract_text_from_file(dest, filename, mime_type)
        if text and text.strip():
            return text.strip()
        return "[ERROR]: Files are images without readable text."
    finally:
        if os.path.exists(dest):
            os.remove(dest)


def refresh_token_if_needed(tokens_path):
    """Check if the access token is expired and refresh if needed."""
    with open(tokens_path, 'r') as f:
        tokens = json.load(f)

    import time
    expiry = tokens.get('expiry_date', 0)
    # expiry_date is in milliseconds
    if expiry / 1000 > time.time() + 60:
        return tokens['access_token']

    # Token expired — try to refresh
    refresh_token = tokens.get('refresh_token')
    if not refresh_token:
        log("  No refresh token available")
        return tokens['access_token']

    # Read client credentials
    creds_path = os.path.expanduser('~/.config/google-drive-creds.json')
    if not os.path.exists(creds_path):
        log("  Client credentials not found, using existing token")
        return tokens['access_token']

    with open(creds_path, 'r') as f:
        creds = json.load(f)

    # Handle both installed and web credential types
    client_info = creds.get('installed') or creds.get('web') or {}
    client_id = client_info.get('client_id', '')
    client_secret = client_info.get('client_secret', '')

    if not client_id or not client_secret:
        log("  Missing client_id/client_secret, using existing token")
        return tokens['access_token']

    # Refresh the token
    log("  Refreshing expired access token...")
    data = urllib.parse.urlencode({
        'client_id': client_id,
        'client_secret': client_secret,
        'refresh_token': refresh_token,
        'grant_type': 'refresh_token'
    }).encode('utf-8')
    req = urllib.request.Request('https://oauth2.googleapis.com/token', data=data)
    try:
        with urllib.request.urlopen(req) as resp:
            result = json.loads(resp.read())
            new_token = result.get('access_token')
            if new_token:
                # Update tokens file
                tokens['access_token'] = new_token
                tokens['expiry_date'] = int(time.time() * 1000) + result.get('expires_in', 3600) * 1000
                with open(tokens_path, 'w') as f:
                    json.dump(tokens, f, indent=2)
                log("  Token refreshed successfully")
                return new_token
    except Exception as e:
        log(f"  Token refresh failed: {e}")

    return tokens['access_token']


def main():
    if len(sys.argv) < 4:
        print("Usage: extract_brochure.py <url_type> <url_or_id> <working_dir> [tokens_path]",
              file=sys.stderr)
        sys.exit(1)

    url_type = sys.argv[1]
    url_or_id = sys.argv[2]
    working_dir = sys.argv[3]
    tokens_path = sys.argv[4] if len(sys.argv) > 4 else os.path.expanduser(
        '~/.config/google-drive-mcp/tokens.json')

    os.makedirs(working_dir, exist_ok=True)

    # Get access token
    access_token = refresh_token_if_needed(tokens_path)

    if url_type == 'GDRIVE_FILE':
        result = process_gdrive_file(url_or_id, access_token, working_dir)
    elif url_type == 'GDRIVE_FOLDER':
        folder_id = extract_folder_id_from_url(url_or_id) if '/' in url_or_id else url_or_id
        if not folder_id:
            result = "[ERROR]: Could not extract folder ID from URL."
        else:
            result = process_gdrive_folder(folder_id, access_token, working_dir)
    elif url_type == 'GDOCS':
        doc_id = extract_doc_id_from_url(url_or_id) if '/' in url_or_id else url_or_id
        if not doc_id:
            result = "[ERROR]: Could not extract document ID from URL."
        else:
            result = process_gdocs(doc_id, access_token, working_dir)
    else:
        result = f"[ERROR]: Unknown URL type: {url_type}"

    # Truncate if needed
    max_len = 95000
    if len(result) > max_len:
        # Find last complete paragraph before limit
        truncated = result[:max_len]
        last_para = truncated.rfind('\n\n')
        if last_para > max_len * 0.8:
            truncated = truncated[:last_para]
        result = truncated + '\n\n[Content truncated at 95,000 character limit]'

    print(result)


if __name__ == '__main__':
    main()
